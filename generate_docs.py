"""
Mechatronics Degree Pack Generator
Creates:
- DOCX: Program Handbook, Lab Manual, Sample Exams, Portfolio Roadmap
- PDF: Same 4 docs (simple, clean layout)
- XLSX: Academic Calendar with 8 semester sheets

Requirements:
  pip install python-docx reportlab openpyxl
Run:
  python generate_docs.py
Output:
  ./out/
"""

from __future__ import annotations

import os
from dataclasses import dataclass
from datetime import date, timedelta
from typing import List, Tuple

# DOCX
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF
from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# XLSX
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter


OUT_DIR = "out"


@dataclass
class DocSpec:
    filename_stem: str
    title: str
    sections: List[Tuple[str, List[str] | List[Tuple[str, List[str]]]]]
    # sections: list of (heading, content)
    # content either list[str] paragraphs OR list of (subheading, list[str])


def ensure_out():
    os.makedirs(OUT_DIR, exist_ok=True)


def _docx_set_base_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for lvl in range(1, 4):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Calibri"
        h.font.bold = True


def _docx_add_title(doc: Document, title: str, subtitle: str | None = None):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(11)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer


def _docx_add_bullets(doc: Document, bullets: List[str]):
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def _docx_add_numbered(doc: Document, items: List[str]):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def build_docx(spec: DocSpec):
    doc = Document()
    _docx_set_base_styles(doc)
    _docx_add_title(doc, spec.title, "Editable DOCX (generated)")

    for heading, content in spec.sections:
        doc.add_heading(heading, level=1)

        if len(content) == 0:
            doc.add_paragraph("")
            continue

        # content is either list[str] or list[tuple]
        if isinstance(content[0], tuple):  # type: ignore[index]
            for sub, paras in content:  # type: ignore[misc]
                doc.add_heading(sub, level=2)
                for para in paras:
                    if para.startswith("• "):
                        _docx_add_bullets(doc, [para[2:]])
                    elif para.startswith("1) ") or para.startswith("1. "):
                        doc.add_paragraph(para)
                    else:
                        doc.add_paragraph(para)
                doc.add_paragraph("")
        else:
            for para in content:  # type: ignore[assignment]
                if para.startswith("BULLETS:"):
                    bullets = [x.strip() for x in para.replace("BULLETS:", "").split("|") if x.strip()]
                    _docx_add_bullets(doc, bullets)
                elif para.startswith("NUMBERED:"):
                    items = [x.strip() for x in para.replace("NUMBERED:", "").split("|") if x.strip()]
                    _docx_add_numbered(doc, items)
                else:
                    doc.add_paragraph(para)

        doc.add_paragraph("")

    path = os.path.join(OUT_DIR, f"{spec.filename_stem}.docx")
    doc.save(path)
    print(f"Wrote {path}")


def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleBig", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=20, leading=24, spaceAfter=12))
    styles.add(ParagraphStyle(name="H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=18, spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, spaceBefore=8, spaceAfter=4))
    styles.add(ParagraphStyle(name="Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=14))
    return styles


def build_pdf(spec: DocSpec):
    styles = _pdf_styles()
    path = os.path.join(OUT_DIR, f"{spec.filename_stem}.pdf")
    doc = SimpleDocTemplate(path, pagesize=LETTER, leftMargin=0.85*inch, rightMargin=0.85*inch, topMargin=0.85*inch, bottomMargin=0.85*inch)

    story = []
    story.append(Paragraph(spec.title, styles["TitleBig"]))
    story.append(Paragraph("Generated PDF (clean print layout)", styles["Body"]))
    story.append(Spacer(1, 10))

    for heading, content in spec.sections:
        story.append(Paragraph(heading, styles["H1"]))

        if len(content) == 0:
            story.append(Spacer(1, 6))
            continue

        if isinstance(content[0], tuple):  # type: ignore[index]
            for sub, paras in content:  # type: ignore[misc]
                story.append(Paragraph(sub, styles["H2"]))
                for para in paras:
                    if para.startswith("BULLETS:"):
                        bullets = [x.strip() for x in para.replace("BULLETS:", "").split("|") if x.strip()]
                        lf = ListFlowable(
                            [ListItem(Paragraph(b, styles["Body"])) for b in bullets],
                            bulletType="bullet",
                            leftIndent=14,
                        )
                        story.append(lf)
                    else:
                        story.append(Paragraph(para, styles["Body"]))
                story.append(Spacer(1, 6))
        else:
            for para in content:  # type: ignore[assignment]
                if para.startswith("BULLETS:"):
                    bullets = [x.strip() for x in para.replace("BULLETS:", "").split("|") if x.strip()]
                    lf = ListFlowable(
                        [ListItem(Paragraph(b, styles["Body"])) for b in bullets],
                        bulletType="bullet",
                        leftIndent=14,
                    )
                    story.append(lf)
                else:
                    story.append(Paragraph(para, styles["Body"]))
            story.append(Spacer(1, 6))

    doc.build(story)
    print(f"Wrote {path}")


def build_calendar_xlsx():
    wb = Workbook()
    wb.remove(wb.active)

    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True)
    bold = Font(bold=True)

    semesters = [
        "Semester 1", "Semester 2", "Semester 3", "Semester 4",
        "Semester 5", "Semester 6", "Semester 7", "Semester 8",
    ]

    for sem in semesters:
        ws = wb.create_sheet(sem)
        ws["A1"] = "Set semester start date (YYYY-MM-DD) in B1:"
        ws["A1"].font = bold
        ws["B1"] = "2026-09-01"

        # header
        headers = ["Week #", "Week Start", "Week End", "Lecture Topics", "Labs/Studios", "Assessments Due"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        # weeks
        for w in range(1, 15):
            r = 3 + w
            ws.cell(row=r, column=1, value=w)
            # Excel formulas for dates
            ws.cell(row=r, column=2, value=f'=DATEVALUE($B$1)+({w-1})*7')
            ws.cell(row=r, column=3, value=f'=DATEVALUE($B$1)+({w-1})*7+6')
            ws.cell(row=r, column=4, value="")
            ws.cell(row=r, column=5, value="")
            ws.cell(row=r, column=6, value="")

        # formatting
        ws.column_dimensions["A"].width = 8
        ws.column_dimensions["B"].width = 14
        ws.column_dimensions["C"].width = 14
        ws.column_dimensions["D"].width = 42
        ws.column_dimensions["E"].width = 30
        ws.column_dimensions["F"].width = 22

        for row in range(4, 18):
            for col in range(1, 7):
                ws.cell(row=row, column=col).alignment = Alignment(vertical="top", wrap_text=True)

    path = os.path.join(OUT_DIR, "Mechatronics_Academic_Calendar.xlsx")
    wb.save(path)
    print(f"Wrote {path}")


def build_specs() -> List[DocSpec]:
    # --- PROGRAM HANDBOOK ---
    handbook = DocSpec(
        filename_stem="Mechatronics_Program_Handbook",
        title="Mechatronics Engineering (B.Eng.) — Program Handbook",
        sections=[
            ("Program Overview", [
                "This program prepares engineers to design, build, test, and maintain integrated mechanical–electrical–software systems.",
                "Hands-on learning is embedded every semester using Arduino/ESP32, Raspberry Pi, sensors, actuators, CAD, and fabrication.",
            ]),
            ("Program Learning Outcomes", [
                "BULLETS: Apply mathematics, physics, and engineering science to mechatronic systems"
                "| Design mechanical components and integrate actuators/sensors"
                "| Build embedded firmware using robust architectures"
                "| Model and tune feedback control systems (PID and beyond)"
                "| Implement robotics perception and autonomy basics (vision, odometry)"
                "| Communicate effectively with technical documentation"
                "| Apply safety, reliability, and professional ethics",
            ]),
            ("Program Structure (8 Semesters)", [
                ("Year 1 — Foundations + Build Skills", [
                    "Semester 1: Mathematics I; Physics I (Mechanics); Programming (Python + C); CAD I; Mechatronics Lab I (Sensors).",
                    "Semester 2: Mathematics II; Physics II (E&M); Circuits I; CAD II; Mechatronics Lab II (Actuators/Mechanisms).",
                ]),
                ("Year 2 — Embedded + Control Foundations", [
                    "Semester 3: Mathematics III (Signals); Statics/Strength; Circuits II; Instrumentation; Embedded Systems I.",
                    "Semester 4: Dynamics; Thermofluids; Control Systems I; Digital Electronics; Integration Studio I.",
                ]),
                ("Year 3 — Robotics + Automation + Vision", [
                    "Semester 5: Robotics I; Control Systems II; Embedded II (RTOS); Mechanical Design; Robotics Lab.",
                    "Semester 6: Industrial Automation; Systems Design; Machine Vision; Energy Conversion; Automation Studio II.",
                ]),
                ("Year 4 — Product-Level Integration + Capstone", [
                    "Semester 7: Advanced Mechatronics; Robotics II; Real-Time Networking; Project Management; Capstone I.",
                    "Semester 8: Safety/Standards/Reliability; Elective; Capstone II.",
                ]),
            ]),
            ("Standard Assessment Model (Hands-on Courses)", [
                "BULLETS: Labs + checkouts (25–35%)"
                "| Lab notebook + documentation (15–25%)"
                "| Practical exams (15–25%)"
                "| Projects (25–40%)",
                "Project rubric (recommended): Functionality 30%, Engineering quality 20%, Software quality 20%, Documentation 15%, Design reasoning 15%.",
            ]),
            ("Course Syllabi (Descriptions + Weekly Topics + Print/Video Notes)", [
                ("MEC110 Mechatronics Lab I: Sensors & Measurement (Arduino)", [
                    "Description: Measurement, sensor interfacing, calibration, uncertainty, filtering, serial logging.",
                    "Learning outcomes: Build reliable sensor circuits; calibrate sensors; log and analyze data; document results.",
                    "Weekly topics (14): 1 Safety + multimeter | 2 Digital IO + debounce | 3 ADC scaling | 4 Thermistor calibration | 5 Ultrasonic timing | 6 I2C sensor + OLED | 7 SPI/SD optional | 8 RC filtering | 9 Sampling/aliasing | 10 Digital filters | 11 Serial CSV + Python plotting | 12 Integration | 13 Project build | 14 Demos + practical.",
                    "Print notes: lab sheets + calibration worksheet + curated reading excerpts.",
                    "Video notes: multimeter basics; breadboard reliability; serial plotting tutorial; sensor calibration walkthrough.",
                    "Assessments: lab checkouts 25%; notebook 20%; practical 25%; mini-project 30%.",
                ]),
                ("MEC120 Mechatronics Lab II: Actuators & Mechanisms (Arduino)", [
                    "Description: PWM, motor drivers, safe wiring, mechanisms, limit switches, homing, basic closed-loop control.",
                    "Weekly topics (14): 1 Power safety | 2 Servo control | 3 DC motor + driver + PWM | 4 Speed measurement | 5 Stepper + driver | 6 Limit switches + homing | 7 Mechanism build | 8 Motion profiles | 9 Stall awareness | 10 Closed-loop speed PI | 11 Backlash | 12 Harness standards | 13 Project build | 14 Demos + practical.",
                    "Print notes: motor driver datasheets + wiring standards + mechanism notes.",
                    "Video notes: PWM fundamentals; motor driver bring-up; stepper homing demo.",
                ]),
                ("EMB210 Embedded Systems I (Arduino/ESP32)", [
                    "Description: Robust embedded firmware patterns: timers, interrupts, comms, fault handling, structured architecture.",
                    "Weekly topics (14): 1 Workflow+Git | 2 Non-blocking state machines | 3 Interrupts | 4 ADC practices | 5 PWM | 6 UART robustness | 7 I2C/SPI debug | 8 Watchdogs | 9 Power/decoupling | 10 Logging | 11 Wi-Fi+MQTT (ESP32) | 12 Architecture | 13 Integration | 14 Practical exam.",
                    "Print notes: embedded style guide + checklists + example driver modules.",
                    "Video notes: interrupts/timers; I2C/SPI debugging; watchdog patterns; MQTT telemetry quickstart.",
                ]),
                ("CON240 Control Systems I (Modeling + PID)", [
                    "Description: Modeling, transfer functions, time response, stability basics, PID tuning and implementation tradeoffs.",
                    "Weekly topics (14): 1 Modeling | 2 Laplace/TF | 3 1st/2nd order response | 4 Poles/zeros | 5 Block diagrams | 6 PID | 7 Tuning methods | 8 Noise/filters | 9 Saturation/anti-windup | 10 Frequency intuition | 11 Root locus intro | 12 Specs/tradeoffs | 13 Project tuning | 14 Review.",
                    "Print notes: control text chapters + tuning worksheets.",
                    "Video notes: PID intuition series; stability basics; tuning walkthrough with logs.",
                ]),
                ("ROB315 Robotics Lab (Raspberry Pi + MCU)", [
                    "Description: Mobile robot integration using Raspberry Pi as high-level compute and MCU for real-time motor control.",
                    "Weekly topics: setup + camera + serial link; encoder odometry calibration; IMU fusion; motor closed-loop; OpenCV pipeline; integration sprints; final autonomy challenge.",
                    "Print notes: robot build guide + wiring harness standard + testing checklist.",
                    "Video notes: Pi camera + OpenCV intro; odometry calibration; IMU filter demonstration.",
                ]),
                ("AUT365 Automation Studio II (Vision Sorting Cell)", [
                    "Description: Mini factory cell using conveyor, vision classification, actuator diverter, fault states, logging.",
                    "Weekly topics: conveyor bring-up; lighting; thresholding/contours; classification; diverter timing; jam detection; state machine; acceptance test plan; demo.",
                    "Print notes: safety/FMEA template + acceptance test checklist.",
                    "Video notes: OpenCV classification pipeline; industrial fault handling patterns.",
                ]),
            ]),
            ("Minimum Kit & Tooling Standard", [
                "BULLETS: Arduino Uno/Nano or compatible"
                "| ESP32 dev board"
                "| Raspberry Pi (3/4/5) + camera"
                "| Sensors: temp/humidity, ultrasonic, IMU, light, encoders"
                "| Actuators: servos, DC motors + driver, stepper + driver"
                "| Multimeter + soldering tools + wire tools"
                "| Optional: 3D printer access; calipers; fastener kit",
            ]),
            ("Safety and Professional Practice", [
                "All labs require: proper power isolation, strain relief, safe wiring, and documented test evidence.",
                "Students must follow academic integrity rules and cite external code/resources.",
            ]),
        ],
    )

    # --- LAB MANUAL ---
    lab_manual = DocSpec(
        filename_stem="Mechatronics_Lab_Manual",
        title="Mechatronics Engineering — Lab Manual",
        sections=[
            ("Universal Lab Sheet Template", [
                "Use this template for every lab.",
                "BULLETS: Lab title, objective, background"
                "| Parts/tools + wiring diagram + pin table"
                "| Procedure with expected observations"
                "| Data tables + required plots"
                "| Checkouts (TA verifies milestones)"
                "| Post-lab questions"
                "| Submission checklist (repo, video, report)"
                "| Safety notes",
            ]),
            ("MEC110 — Lab Sheets (Sensors & Measurement)", [
                ("Lab 1: Digital IO + Debounce", [
                    "Objective: Build a stable button input and demonstrate debounce.",
                    "Parts: Arduino, button, resistor (or internal pull-up), LED, breadboard, wires.",
                    "Procedure: wire pull-up; read raw; observe bounce; implement debounce; verify stable count.",
                    "Data: raw vs debounced press timestamps; count accuracy over 30 presses.",
                    "Checkouts: explain pull-up; show stable count; show code is non-blocking.",
                ]),
                ("Lab 3: ADC Scaling + Noise", [
                    "Objective: Convert ADC counts to volts and characterize noise.",
                    "Procedure: read 1000 samples from pot; compute volts; histogram; mean/std dev; discuss resolution.",
                    "Checkouts: explain reference voltage; show correct scaling formula.",
                ]),
                ("Lab 4: Thermistor Calibration", [
                    "Objective: Calibrate thermistor with 3 reference points; compute curve fit; estimate error.",
                    "Procedure: measure resistance vs temp (ice/room/warm); fit simplified model; validate on a 4th point.",
                    "Submission: plot + error table + calibration constants.",
                ]),
                ("Lab 11: Serial CSV + Python Plot", [
                    "Objective: stream CSV; plot live; save dataset; compute summary stats.",
                    "Procedure: Arduino prints CSV; Python reads serial, plots; saves file; computes mean/std/peak-to-peak.",
                    "Checkouts: live plot + saved dataset + short analysis paragraph.",
                ]),
            ]),
            ("MEC120 — Lab Sheets (Actuators & Mechanisms)", [
                ("Lab A: DC Motor PWM with Driver", [
                    "Objective: Drive DC motor safely with PWM; verify speed steps and thermal safety.",
                    "Safety: separate motor supply, common ground, never power motor from MCU 5V.",
                    "Procedure: wire TB6612FNG; command PWM steps; record speed proxy; inspect heating.",
                    "Checkouts: demonstrate stop on fault + safe wiring photo.",
                ]),
                ("Lab B: Stepper Homing with Limit Switch", [
                    "Objective: Implement homing routine; test repeatability (10 cycles).",
                    "Procedure: drive stepper; detect switch; back off; re-approach slowly; record home position error.",
                    "Submission: repeatability plot; discussion of switch bounce mitigation.",
                ]),
                ("Lab C: Build a Linear Stage", [
                    "Objective: Assemble belt or lead-screw stage; measure backlash and positioning error.",
                    "Procedure: build; command moves; measure with ruler/calipers; compute error stats.",
                ]),
            ]),
            ("EMB210 — Embedded Lab Sheets (Robust Firmware)", [
                "BULLETS: Non-blocking scheduler"
                "| Interrupt timing measurement"
                "| I2C/SPI debug checklist"
                "| Structured telemetry + CRC"
                "| Watchdog + fault handling states"
            ]),
            ("CON240 — PID Implementation Labs", [
                "BULLETS: Step response identification"
                "| PID tuning with logging plots"
                "| Saturation + anti-windup"
                "| Load disturbance rejection"
            ]),
            ("ROB315 — Robotics Lab Sheets (Pi + MCU)", [
                "BULLETS: Pi bring-up + camera"
                "| Encoder odometry calibration"
                "| IMU complementary filter"
                "| Closed-loop motor speed control on MCU"
                "| OpenCV line follow / object tracking"
                "| Final integration challenge"
            ]),
            ("AUT365 — Automation Studio Lab Sheets", [
                "BULLETS: Conveyor bring-up + speed control"
                "| Lighting + threshold pipeline"
                "| Classification + diverter timing"
                "| Fault states (jam timeout, sensor failure)"
                "| Logging + acceptance tests"
            ]),
        ],
    )

    # --- EXAMS ---
    exams = DocSpec(
        filename_stem="Mechatronics_Sample_Exams",
        title="Mechatronics Engineering — Sample Exams Pack",
        sections=[
            ("Circuits I — Sample Midterm (90 minutes)", [
                "BULLETS: KCL/KVL node voltage problem (20)"
                "| Thevenin equivalent (15)"
                "| RC transient Vc(t), time to 90% (20)"
                "| AC basics: impedance and phase (15)"
                "| LED current limiting design (10)"
                "| Troubleshooting from symptoms (20)",
                "Marking: show method; correct units; interpret physical reasonableness.",
            ]),
            ("Control Systems I — Sample Final (2 hours)", [
                "BULLETS: Stability from poles/zeros (20)"
                "| Estimate parameters from step response plot (20)"
                "| PID effects on overshoot/steady-state/noise (20)"
                "| Integrator windup + anti-windup method (20)"
                "| Design PID to meet specs (20)",
                "Allow partial credit for correct reasoning even if tuning is imperfect.",
            ]),
            ("Embedded Systems I — Sample Practical (75 minutes)", [
                "Tasks:",
                "BULLETS: Implement non-blocking periodic sensor read + rolling average"
                "| Add fault detection and safe behavior"
                "| Stream structured telemetry and show a plot/log"
                "| Short viva: explain architecture and timing",
                "Rubric: functionality 40%, robustness 20%, code quality 20%, documentation 10%, explanation 10%.",
            ]),
            ("Robotics Lab — Final Challenge", [
                "Robot must follow a path, avoid an obstacle, and stop at a target marker.",
                "Scoring: stability, repeatability, robustness to lighting changes, logging quality.",
            ]),
            ("Automation Studio — Demo + Oral Defense", [
                "Must demonstrate classification accuracy, jam detection, fail-safe state machine, and logs.",
                "Oral questions: hazards, mitigations, verification plan.",
            ]),
        ],
    )

    # --- PORTFOLIO ---
    portfolio = DocSpec(
        filename_stem="Mechatronics_Portfolio_Roadmap",
        title="Mechatronics Engineering — 8-Semester Portfolio Roadmap",
        sections=[
            ("Portfolio Standard (Every Artifact)", [
                "BULLETS: Git repo with code + CAD + schematics + BOM"
                "| Demo video (1–3 minutes)"
                "| Test evidence (plots, logs, checklists)"
                "| Short report (problem, design, results, limitations, next steps)",
            ]),
            ("Semester-by-Semester Artifacts", [
                ("Semester 1", [
                    "BULLETS: Sensor Calibration Notebook (plots + uncertainty notes)"
                    "| Smart Desk Monitor (Arduino: light/temp + alert)",
                ]),
                ("Semester 2", [
                    "BULLETS: Mini Linear Stage or Conveyor (CAD + mechanism + demo)"
                    "| Motor Control Safety Build (power + grounding + harness standards)",
                ]),
                ("Semester 3", [
                    "BULLETS: Wireless Condition Monitor (ESP32 + dashboard)"
                    "| Reusable Driver Library Pack (I2C/SPI/UART modules)",
                ]),
                ("Semester 4", [
                    "BULLETS: PID Speed Controller with Encoder (tuning plots)"
                    "| Bring-up Guide (reproducible build + debug steps)",
                ]),
                ("Semester 5", [
                    "BULLETS: Autonomous Rover v1 (Pi + MCU: odometry + IMU fusion)"
                    "| Vision Demo (line follow or object tracking)",
                ]),
                ("Semester 6", [
                    "BULLETS: Vision Sorting Cell (fault handling + logs)"
                    "| Safety + FMEA Case Study (hazard table + mitigations + tests)",
                ]),
                ("Semester 7", [
                    "BULLETS: Capstone Proposal Package (requirements, architecture, BOM, schedule)"
                    "| Prototype Subsystem with validation evidence",
                ]),
                ("Semester 8", [
                    "BULLETS: Capstone Final Product (demo + validation report + repo + poster)"
                    "| Engineering Retrospective (tradeoffs, failures, fixes, next steps)",
                ]),
            ]),
            ("Capstone Gold Standard Checklist", [
                "BULLETS: Requirements + acceptance tests"
                "| CAD drawings + assembly"
                "| Schematics + wiring harness diagram"
                "| Firmware architecture + fault handling"
                "| Verification matrix (req → tests)"
                "| Logs/plots + validation report"
                "| Costed BOM + procurement notes"
                "| Safety/risk notes + mitigations",
            ]),
        ],
    )

    return [handbook, lab_manual, exams, portfolio]


def main():
    ensure_out()
    specs = build_specs()

    for s in specs:
        build_docx(s)
        build_pdf(s)

    build_calendar_xlsx()
    print("\nAll files generated in ./out/\n")


if __name__ == "__main__":
    main()
